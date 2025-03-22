#### md_to_word.py
####
#### This script converts a markdown file to a Word document.
#### It uses the python-docx library to create and manipulate Word documents.
#### The script reads the markdown content, parses headings, paragraphs, images and other formatting elements,
#### and converts them to their corresponding Word document equivalents.


import os
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


def convert_md_to_docx(md_file_path, output_dir, image_dir='visualizations'):
    """Convert markdown report to Word document"""
    # Create paths
    docx_path = os.path.join(output_dir, 'executive_report.docx')
    full_image_dir = os.path.join(output_dir, image_dir)

    # Create new document
    doc = Document()

    # Configure default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Read markdown content
    try:
        with open(md_file_path, 'r') as f:
            content = f.read()
    except FileNotFoundError:
        print(f"Error: Markdown file not found at {md_file_path}")
        return False

    # Split into sections
    sections = re.split(r'\n## ', content)

    # Process main header
    main_header = sections[0].replace('# ', '').strip()
    doc.add_heading(main_header, 0)

    for section in sections[1:]:
        lines = section.split('\n')
        section_title = lines[0].strip()

        # Add section heading
        doc.add_heading(section_title, level=1)

        # Process section content
        for line in lines[1:]:
            line = line.strip()

            # Handle images
            if line.startswith('!['):
                alt_text, image_path = re.match(r'!\[(.*?)\]\((.*?)\)', line).groups()
                full_image_path = os.path.join(output_dir, image_path)

                if os.path.exists(full_image_path):
                    para = doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(full_image_path, width=Inches(6))
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    doc.add_paragraph(alt_text, style='Caption')

            # Handle bullet points
            elif line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:].strip())

            # Handle subheadings
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)

            # Handle bold text
            elif '**' in line:
                p = doc.add_paragraph()
                for part in re.split(r'(\*\*.*?\*\*)', line):
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

            # Regular paragraphs
            elif line:
                doc.add_paragraph(line)

        # Add section break
        doc.add_page_break()

    # Save document
    try:
        doc.save(docx_path)
        print(f"DOCX report saved to '{docx_path}'")
        return True
    except Exception as e:
        print(f"Error saving DOCX file: {str(e)}")
        return False